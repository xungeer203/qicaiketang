import docx
from docx.shared import Mm, Pt
import set_cell_border as cb
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# todo: 表格跨页不要被打断，这是个参考：Table.allow_break_across_pages

def export2word(pinyin_group, zi_group, char_size, char_size_half, title="看拼音写词语", file_name="text.docx"):
    # title = "第一单元"
    doc = docx.Document()
    # doc.add_heading(title, 0)
    # doc.add_paragraph(title)

    doc.add_paragraph()
    run = doc.paragraphs[0].add_run(title)
    run.font.name = "楷体"
    run.font.size = Pt(10)
    r = run._element
    r.rPr.rFonts.set(qn("w:eastAsia"), "楷体")  # refer to "note1"



    for pinyins, zis in zip(pinyin_group, zi_group):
        widths = [Mm(char_size) if pinyin != "" else Mm(char_size_half) for pinyin in pinyins]
        table = doc.add_table(rows=2, cols=len(pinyins))

        row = table.rows[0].cells
        for cell, pinyin in zip(row, pinyins):
            cell.text = pinyin
            cell_font = cell.paragraphs[0].runs[0].font
            cell_font.size = Pt(10) if len(pinyin) < 6 else Pt(8)
            cell_paragraph_format = cell.paragraphs[0].paragraph_format
            cell_paragraph_format.space_after = Pt(0)
            cell_paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER


        row = table.rows[1].cells
        for cell, zi in zip(row, zis):
            text = ""  # todo: 哪些字需要打印出来？
            # text = zi
            # cell.text = text
            # cell_font.size = Pt(8)
            # cell_font.name = 'SimSun-ExtB'  # 英语这样输入，改字号，改字体

            # input text
            run = cell.paragraphs[0].add_run(text)  # note1：中文需要这样的输入，后续才能改字体。　ref: https://blog.csdn.net/qq_27017791/article/details/108897521, clipped in Onenote
            run.font.name = "楷体"
            run.font.size = Pt(11)
            r = run._element
            r.rPr.rFonts.set(qn("w:eastAsia"), "楷体")

            # cell boader
            if zi != "":
                cb.set_cell_border(cell,
                                   bottom={"sz": 3, "val": "single", "color": "#FF0000", "space": "0"}
                                   )  # refer to instructions of this function.

            cell_paragraph_format = cell.paragraphs[0].paragraph_format
            cell_paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell_paragraph_format.space_after = Pt(0)

        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width

        for row in table.rows:
            # for idx, width in enumerate(widths):
            #     # row.cells[idx].width = width
            row.height = Mm(7)
        # table.rows[0].cells[0].height = Mm(10)

        run = doc.add_paragraph("")
        run.paragraph_format.space_after = Pt(0)









    doc.save(file_name)
